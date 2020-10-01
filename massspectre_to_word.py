from docx import Document
from docx.shared import Inches
import os
from docx.shared import Pt

#папка с картинками
picture_dir = '/example_path'
#получаем список файлов и сортируем по дате
os.chdir(picture_dir)
files = filter(os.path.isfile, os.listdir(picture_dir))
files = [os.path.join(picture_dir, f) for f in files]
files.sort(key=lambda x: os.path.getmtime(x))

document = Document()
p = document.add_paragraph()
r = p.add_run()
r.font.size = Pt(14)
for file in files:
    print('insert ==> ', file)
    r.add_picture(file,  width=Inches(6.0))
    r.add_text(file[len(picture_dir):len(file) - 3])

document.save('result.docx')